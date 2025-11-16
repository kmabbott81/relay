from __future__ import annotations

import csv
import hashlib
import json
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    import yaml  # type: ignore
except Exception as e:
    raise RuntimeError("PyYAML required. Install with: pip install pyyaml") from e

try:
    import jsonschema  # type: ignore
except Exception as e:
    raise RuntimeError("jsonschema required. Install with: pip install jsonschema") from e

from docx import Document
from docx.shared import Pt
from jinja2 import Environment, StrictUndefined, Template, TemplateError
from jinja2.sandbox import SandboxedEnvironment

TEMPLATES_DIR = Path("templates")
CUSTOM_TEMPLATES_DIR = Path("templates/custom")
OUTPUT_DIR = Path("runs/ui/templates")
STYLES_DIR = Path("styles")
SCHEMA_PATH = Path("schemas/template.json")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
STYLES_DIR.mkdir(parents=True, exist_ok=True)

# Template compilation cache
# Key: (template_path, content_hash) -> Value: compiled Jinja2 Template
_template_cache: dict[tuple[str, str], Template] = {}
_cache_stats = {"hits": 0, "misses": 0}


class TemplateValidationError(Exception):
    """Raised when template validation fails with user-friendly message."""

    def __init__(self, template_path: str, errors: list[str]):
        self.template_path = template_path
        self.errors = errors
        message = f"Template validation failed for {template_path}:\n" + "\n".join(f"  - {e}" for e in errors)
        super().__init__(message)


class TemplateRenderError(Exception):
    """Raised when template rendering fails with user-friendly message."""

    pass


@dataclass
class InputDef:
    """Definition of a template input field."""

    id: str
    label: str
    type: str
    required: bool = False
    default: Any = None
    help: str = ""
    placeholder: str = ""
    validators: dict[str, Any] = field(default_factory=dict)


@dataclass
class TemplateDef:
    """Complete template definition with metadata and inputs."""

    path: Path
    name: str
    version: str
    description: str
    context: str
    inputs: list[InputDef]
    body: str
    style: str | None = None

    @property
    def key(self) -> str:
        """Generate key from filename."""
        return self.path.stem


def _load_schema() -> dict[str, Any]:
    """Load the template JSON schema."""
    if not SCHEMA_PATH.exists():
        raise FileNotFoundError(f"Template schema not found at {SCHEMA_PATH}")
    return json.loads(SCHEMA_PATH.read_text(encoding="utf-8"))


def _validate_template_data(data: dict[str, Any], template_path: str) -> None:
    """
    Validate template data against JSON schema.

    Args:
        data: Template YAML data
        template_path: Path to template file for error messages

    Raises:
        TemplateValidationError: If validation fails
    """
    schema = _load_schema()
    validator = jsonschema.Draft7Validator(schema)
    errors = []

    for error in validator.iter_errors(data):
        # Convert jsonschema errors to friendly messages
        field_path = ".".join(str(p) for p in error.path) if error.path else "root"
        errors.append(f"{field_path}: {error.message}")

    if errors:
        raise TemplateValidationError(template_path, errors)


def _load_yaml(path: Path) -> dict[str, Any]:
    """Load and parse YAML file."""
    try:
        return yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    except yaml.YAMLError as e:
        raise TemplateValidationError(str(path), [f"Invalid YAML syntax: {e}"]) from e


def _parse_template(path: Path) -> TemplateDef:
    """
    Load and validate a template from YAML file.

    Args:
        path: Path to template YAML file

    Returns:
        TemplateDef object

    Raises:
        TemplateValidationError: If template is invalid
    """
    data = _load_yaml(path)

    # Validate against schema
    _validate_template_data(data, str(path))

    # Parse inputs
    inputs = []
    for inp_data in data.get("inputs", []):
        inputs.append(
            InputDef(
                id=inp_data["id"],
                label=inp_data["label"],
                type=inp_data["type"],
                required=inp_data.get("required", False),
                default=inp_data.get("default"),
                help=inp_data.get("help", ""),
                placeholder=inp_data.get("placeholder", ""),
                validators=inp_data.get("validators", {}),
            )
        )

    # Extract rendering body
    rendering = data.get("rendering", {})
    if not rendering.get("body"):
        raise TemplateValidationError(str(path), ["rendering.body is required"])

    return TemplateDef(
        path=path,
        name=data["name"],
        version=data["version"],
        description=data["description"],
        context=data["context"],
        inputs=inputs,
        body=rendering["body"],
        style=data.get("style"),
    )


def list_templates() -> list[TemplateDef]:
    """
    List all valid templates from templates/ directory.

    Returns:
        List of TemplateDef objects. Invalid templates are logged but not returned.
    """
    TEMPLATES_DIR.mkdir(exist_ok=True, parents=True)
    CUSTOM_TEMPLATES_DIR.mkdir(exist_ok=True, parents=True)

    out: list[TemplateDef] = []
    search_paths = [
        *sorted(TEMPLATES_DIR.glob("*.yaml")),
        *sorted(CUSTOM_TEMPLATES_DIR.glob("*.yaml")),
    ]

    for yml in search_paths:
        try:
            template = _parse_template(yml)
            out.append(template)
        except TemplateValidationError as e:
            # Log but don't crash - allow UI to show error
            print(f"Warning: {e}")
        except Exception as e:
            print(f"Warning: Failed to load {yml}: {e}")

    return out


# Custom Jinja2 filters for template safety
def to_slug(text: str) -> str:
    """Convert text to URL-safe slug."""
    text = str(text).lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[-\s]+", "-", text)
    return text


def to_title(text: str) -> str:
    """Convert text to title case."""
    return str(text).title()


def _compute_template_hash(template_body: str) -> str:
    """
    Compute hash of template body for cache invalidation.

    Args:
        template_body: Template body text

    Returns:
        SHA256 hash of template body
    """
    return hashlib.sha256(template_body.encode("utf-8")).hexdigest()[:16]


def _get_cached_template(template: TemplateDef, env: Environment) -> Template:
    """
    Get cached compiled template or compile and cache it.

    Args:
        template: Template definition
        env: Jinja2 environment

    Returns:
        Compiled Jinja2 Template
    """
    cache_key = (str(template.path), _compute_template_hash(template.body))

    if cache_key in _template_cache:
        _cache_stats["hits"] += 1
        return _template_cache[cache_key]

    # Cache miss - compile template
    _cache_stats["misses"] += 1
    compiled = env.from_string(template.body)
    _template_cache[cache_key] = compiled

    return compiled


def get_cache_stats() -> dict[str, int]:
    """
    Get template cache statistics.

    Returns:
        Dictionary with cache hits and misses
    """
    return _cache_stats.copy()


def clear_template_cache() -> None:
    """Clear the template compilation cache."""
    global _template_cache, _cache_stats
    _template_cache.clear()
    _cache_stats = {"hits": 0, "misses": 0}


def _create_sandbox_env(context: str) -> Environment:
    """
    Create a sandboxed Jinja2 environment with safe defaults.

    Args:
        context: Output context ("markdown", "docx", "html")

    Returns:
        Configured SandboxedEnvironment
    """
    # Enable autoescape for HTML/DOCX contexts
    autoescape = context in ("html", "docx")

    env = SandboxedEnvironment(
        autoescape=autoescape,
        undefined=StrictUndefined,
        trim_blocks=True,
        lstrip_blocks=True,
    )

    # Add custom filters
    env.filters["to_slug"] = to_slug
    env.filters["to_title"] = to_title

    # Restrict to safe built-in filters only
    safe_filters = {
        "lower",
        "upper",
        "title",
        "replace",
        "join",
        "length",
        "round",
        "default",
        "safe",
        "escape",
        "e",
        "map",
        "select",
    }

    # Remove any unsafe filters
    for filter_name in list(env.filters.keys()):
        if filter_name not in safe_filters and not filter_name.startswith("to_"):
            del env.filters[filter_name]

    return env


def validate_inputs(template: TemplateDef, values: dict[str, Any]) -> list[str]:
    """
    Validate input values against template input definitions.

    Args:
        template: Template definition
        values: Input values to validate

    Returns:
        List of validation error messages (empty if valid)
    """
    errors = []

    for inp in template.inputs:
        value = values.get(inp.id)

        # Check required
        if inp.required and (value is None or value == ""):
            errors.append(f"{inp.label} is required")
            continue

        # Skip validation if empty and not required
        if value is None or value == "":
            continue

        # Type validation
        if inp.type == "int":
            try:
                val = int(value)
                if "min" in inp.validators and val < inp.validators["min"]:
                    errors.append(f"{inp.label} must be at least {inp.validators['min']}")
                if "max" in inp.validators and val > inp.validators["max"]:
                    errors.append(f"{inp.label} must be at most {inp.validators['max']}")
            except (ValueError, TypeError):
                errors.append(f"{inp.label} must be an integer")

        elif inp.type == "float":
            try:
                val = float(value)
                if "min" in inp.validators and val < inp.validators["min"]:
                    errors.append(f"{inp.label} must be at least {inp.validators['min']}")
                if "max" in inp.validators and val > inp.validators["max"]:
                    errors.append(f"{inp.label} must be at most {inp.validators['max']}")
            except (ValueError, TypeError):
                errors.append(f"{inp.label} must be a number")

        elif inp.type == "email":
            email_pattern = r"^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$"
            if not re.match(email_pattern, str(value)):
                errors.append(f"{inp.label} must be a valid email address")

        elif inp.type == "url":
            url_pattern = r"^https?://.+\..+"
            if not re.match(url_pattern, str(value)):
                errors.append(f"{inp.label} must be a valid URL (http:// or https://)")

        elif inp.type == "enum":
            choices = inp.validators.get("choices", [])
            if value not in choices:
                errors.append(f"{inp.label} must be one of: {', '.join(choices)}")

        elif inp.type == "multiselect":
            choices = inp.validators.get("choices", [])
            if isinstance(value, list):
                for v in value:
                    if v not in choices:
                        errors.append(f"{inp.label}: '{v}' is not a valid choice")

        # String length validation
        if inp.type in ("string", "text") and isinstance(value, str):
            if "min" in inp.validators and len(value) < inp.validators["min"]:
                errors.append(f"{inp.label} must be at least {inp.validators['min']} characters")
            if "max" in inp.validators and len(value) > inp.validators["max"]:
                errors.append(f"{inp.label} must be at most {inp.validators['max']} characters")

        # Regex validation
        if "regex" in inp.validators and isinstance(value, str):
            pattern = inp.validators["regex"]
            if not re.match(pattern, value):
                errors.append(f"{inp.label} does not match required pattern")

    return errors


def render_template(template: TemplateDef, variables: dict[str, Any]) -> str:
    """
    Render a template with sandboxed Jinja2 environment and caching.

    Args:
        template: Template definition
        variables: Variable values for rendering

    Returns:
        Rendered template text

    Raises:
        TemplateRenderError: If rendering fails
    """
    # Validate inputs first
    validation_errors = validate_inputs(template, variables)
    if validation_errors:
        raise TemplateRenderError("Validation failed:\n" + "\n".join(f"  - {e}" for e in validation_errors))

    # Create sandboxed environment
    env = _create_sandbox_env(template.context)

    try:
        # Use cached compiled template
        tmpl = _get_cached_template(template, env)
        return tmpl.render(**variables)
    except TemplateError as e:
        raise TemplateRenderError(f"Template rendering failed: {e}") from e
    except Exception as e:
        raise TemplateRenderError(f"Unexpected error during rendering: {e}") from e


def export_markdown(text: str, fname: str) -> Path:
    """
    Export text as markdown file.

    Args:
        text: Content to export
        fname: Base filename (without extension)

    Returns:
        Path to created file
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    p = OUTPUT_DIR / f"{fname}.md"
    p.write_text(text, encoding="utf-8")
    return p


def export_docx(text: str, fname: str, heading: str = "DJP Output", style_path: str | None = None) -> Path:
    """
    Export text as DOCX file with optional style template.

    Args:
        text: Content to export
        fname: Base filename (without extension)
        heading: Document heading
        style_path: Optional path to base style DOCX

    Returns:
        Path to created file
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load base style if provided
    if style_path and Path(style_path).exists():
        try:
            doc = Document(style_path)
            # Clear existing content but keep styles
            for element in doc.element.body:
                doc.element.body.remove(element)
        except Exception:
            # Fallback to new document if style loading fails
            doc = Document()
    else:
        doc = Document()

    # Add heading
    doc.add_heading(heading, level=1)

    # Add content
    for line in text.splitlines():
        if line.strip():
            para = doc.add_paragraph(line)
            for run in para.runs:
                run.font.size = Pt(11)

    p = OUTPUT_DIR / f"{fname}.docx"
    doc.save(p)
    return p


def estimate_template_cost(template: TemplateDef, variables: dict[str, Any]) -> dict[str, Any]:
    """
    Estimate cost of running a template through DJP.

    Args:
        template: Template definition
        variables: Variable values

    Returns:
        Dictionary with cost_usd, tokens_estimated, margin_pct
    """
    # Import costs module
    try:
        from .costs import project_workflow_cost
    except ImportError:
        # Fallback if costs module unavailable
        return {"cost_usd": 0.0, "tokens_estimated": 0, "margin_pct": 50.0, "note": "Cost estimation unavailable"}

    # Render template to estimate length
    try:
        rendered = render_template(template, variables)
        prompt_length = len(rendered)
    except Exception:
        # Use template body length as fallback
        prompt_length = len(template.body)

    # Rough token estimation: ~4 chars per token
    estimated_prompt_tokens = prompt_length // 4

    # Project workflow cost (uses default models)
    projection = project_workflow_cost(
        max_debaters=3,
        max_tokens=1200,
        require_citations=0,
        fastpath=False,
    )

    # Add template-specific token overhead
    total_tokens = projection.total_tokens_projected + estimated_prompt_tokens

    # Conservative cost estimate (add 50% margin)
    base_cost = projection.total_cost
    margin_pct = 50.0
    estimated_cost = base_cost * (1 + margin_pct / 100)

    return {
        "cost_usd": round(estimated_cost, 4),
        "tokens_estimated": total_tokens,
        "margin_pct": margin_pct,
        "breakdown": {
            "base_workflow": round(base_cost, 4),
            "template_tokens": estimated_prompt_tokens,
        },
    }


def check_budget(
    estimated_cost: float, estimated_tokens: int, budget_usd: float | None = None, budget_tokens: int | None = None
) -> tuple[bool, str, str]:
    """
    Check if estimated cost/tokens fit within budget.

    Args:
        estimated_cost: Estimated cost in USD
        estimated_tokens: Estimated token usage
        budget_usd: Optional USD budget limit
        budget_tokens: Optional token budget limit

    Returns:
        Tuple of (within_budget, warning_message, error_message)
    """
    warnings = []
    errors = []

    # Check USD budget
    if budget_usd is not None:
        ratio = estimated_cost / budget_usd if budget_usd > 0 else float("inf")
        if ratio > 1.0:
            errors.append(f"Estimated cost ${estimated_cost:.4f} exceeds budget ${budget_usd:.4f}")
        elif ratio > 0.9:
            warnings.append(f"Estimated cost ${estimated_cost:.4f} is {ratio * 100:.0f}% of budget ${budget_usd:.4f}")

    # Check token budget
    if budget_tokens is not None:
        ratio = estimated_tokens / budget_tokens if budget_tokens > 0 else float("inf")
        if ratio > 1.0:
            errors.append(f"Estimated tokens {estimated_tokens:,} exceeds budget {budget_tokens:,}")
        elif ratio > 0.9:
            warnings.append(f"Estimated tokens {estimated_tokens:,} is {ratio * 100:.0f}% of budget {budget_tokens:,}")

    within_budget = len(errors) == 0
    warning_msg = "; ".join(warnings) if warnings else ""
    error_msg = "; ".join(errors) if errors else ""

    return within_budget, warning_msg, error_msg


def create_template_artifact(
    template: TemplateDef,
    variables: dict[str, Any],
    rendered_body: str,
    result: dict[str, Any],
    cost_projection: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """
    Create artifact JSON for template run.

    Args:
        template: Template used
        variables: Input variables
        rendered_body: Rendered template text
        result: DJP result dict
        cost_projection: Optional cost projection data

    Returns:
        Artifact dictionary
    """
    artifact = {
        "template": {
            "name": template.name,
            "version": template.version,
            "key": template.key,
            "context": template.context,
        },
        "inputs": variables,
        "provenance": {
            "template_body": rendered_body,
            "resolved_inputs": variables,
            "timestamp": int(time.time()),
        },
        "result": result,
    }

    if cost_projection:
        artifact["cost_projection"] = cost_projection

    return artifact


def load_csv_for_batch(csv_path: str | Path, template: TemplateDef) -> tuple[list[dict[str, Any]], list[str]]:
    """
    Load CSV file and validate rows for batch template processing.

    Args:
        csv_path: Path to CSV file
        template: Template definition

    Returns:
        Tuple of (valid_rows, validation_errors)
    """
    csv_path = Path(csv_path)
    if not csv_path.exists():
        return [], [f"CSV file not found: {csv_path}"]

    # Get required input IDs from template
    required_ids = {inp.id for inp in template.inputs if inp.required}
    all_ids = {inp.id for inp in template.inputs}

    rows = []
    errors = []

    try:
        with open(csv_path, encoding="utf-8") as f:
            reader = csv.DictReader(f)

            if not reader.fieldnames:
                return [], ["CSV file is empty or has no headers"]

            # Check that all required columns are present
            missing = required_ids - set(reader.fieldnames)
            if missing:
                errors.append(f"Missing required columns: {', '.join(missing)}")
                return [], errors

            for row_num, row in enumerate(reader, start=2):  # Start at 2 (1 is header)
                # Extract only template input fields
                row_data = {key: val for key, val in row.items() if key in all_ids}

                # Validate row
                row_errors = validate_inputs(template, row_data)
                if row_errors:
                    errors.append(f"Row {row_num}: {'; '.join(row_errors)}")
                else:
                    rows.append(row_data)

    except csv.Error as e:
        return [], [f"CSV parsing error: {e}"]
    except Exception as e:
        return [], [f"Error reading CSV: {e}"]

    return rows, errors


def estimate_batch_cost(template: TemplateDef, rows: list[dict[str, Any]]) -> dict[str, Any]:
    """
    Estimate total cost for batch processing.

    Args:
        template: Template definition
        rows: List of input variable dictionaries

    Returns:
        Dictionary with total cost estimation
    """
    if not rows:
        return {"total_cost_usd": 0.0, "total_tokens": 0, "num_rows": 0, "per_row_estimates": []}

    per_row_estimates = []
    total_cost = 0.0
    total_tokens = 0

    for row in rows:
        est = estimate_template_cost(template, row)
        per_row_estimates.append(est)
        total_cost += est["cost_usd"]
        total_tokens += est["tokens_estimated"]

    return {
        "total_cost_usd": round(total_cost, 4),
        "total_tokens": total_tokens,
        "num_rows": len(rows),
        "per_row_estimates": per_row_estimates,
    }


def process_batch_dry_run(
    template: TemplateDef, rows: list[dict[str, Any]], budget_usd: float | None = None, budget_tokens: int | None = None
) -> dict[str, Any]:
    """
    Dry run batch processing to show cost projection and budget check.

    Args:
        template: Template definition
        rows: List of input variable dictionaries
        budget_usd: Optional USD budget
        budget_tokens: Optional token budget

    Returns:
        Dictionary with dry run results
    """
    # Estimate costs
    batch_est = estimate_batch_cost(template, rows)

    # Check budget
    within_budget = True
    budget_warnings = []
    budget_errors = []

    if budget_usd or budget_tokens:
        within_budget, warning, error = check_budget(
            batch_est["total_cost_usd"], batch_est["total_tokens"], budget_usd, budget_tokens
        )
        if warning:
            budget_warnings.append(warning)
        if error:
            budget_errors.append(error)

    return {
        "num_rows": batch_est["num_rows"],
        "total_cost_usd": batch_est["total_cost_usd"],
        "total_tokens": batch_est["total_tokens"],
        "within_budget": within_budget,
        "warnings": budget_warnings,
        "errors": budget_errors,
    }


def clone_template(source_template: TemplateDef, new_name: str, new_description: str | None = None) -> Path:
    """
    Clone an existing template to templates/custom/ directory.

    Args:
        source_template: Template to clone
        new_name: Name for the cloned template
        new_description: Optional new description

    Returns:
        Path to created template file

    Raises:
        ValueError: If target already exists or name is invalid
    """
    # Create safe slug for filename
    slug = to_slug(new_name)
    if not slug:
        raise ValueError(f"Invalid template name: {new_name}")

    # Ensure custom directory exists
    CUSTOM_TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    # Check if file already exists
    target_path = CUSTOM_TEMPLATES_DIR / f"{slug}.yaml"
    if target_path.exists():
        raise ValueError(f"Template already exists: {target_path.name}")

    # Load source YAML
    source_data = _load_yaml(source_template.path)

    # Update metadata
    source_data["name"] = new_name
    if new_description:
        source_data["description"] = new_description

    # Increment version (minor bump)
    try:
        major, minor = source_data["version"].split(".")
        source_data["version"] = f"{major}.{int(minor) + 1}"
    except (ValueError, KeyError):
        source_data["version"] = "1.0"

    # Write to custom directory
    target_path.write_text(yaml.dump(source_data, sort_keys=False), encoding="utf-8")

    return target_path


def update_template_yaml(template_path: Path, yaml_content: str) -> tuple[bool, list[str]]:
    """
    Update a template YAML file with validation.

    Args:
        template_path: Path to template file
        yaml_content: New YAML content

    Returns:
        Tuple of (success, validation_errors)
    """
    # Only allow updates to custom templates
    if not str(template_path).startswith(str(CUSTOM_TEMPLATES_DIR)):
        return False, ["Cannot modify built-in templates. Clone to custom/ first."]

    try:
        # Parse YAML
        data = yaml.safe_load(yaml_content)
        if not data:
            return False, ["Empty YAML content"]

        # Validate against schema
        _validate_template_data(data, str(template_path))

        # Write to file
        template_path.write_text(yaml_content, encoding="utf-8")

        return True, []

    except yaml.YAMLError as e:
        return False, [f"Invalid YAML syntax: {e}"]
    except TemplateValidationError as e:
        return False, e.errors
    except Exception as e:
        return False, [f"Error updating template: {e}"]


def delete_custom_template(template_path: Path) -> tuple[bool, str]:
    """
    Delete a custom template file.

    Args:
        template_path: Path to template file

    Returns:
        Tuple of (success, error_message)
    """
    # Only allow deletion of custom templates
    if not str(template_path).startswith(str(CUSTOM_TEMPLATES_DIR)):
        return False, "Cannot delete built-in templates"

    if not template_path.exists():
        return False, "Template file not found"

    try:
        template_path.unlink()
        return True, ""
    except Exception as e:
        return False, f"Error deleting template: {e}"
